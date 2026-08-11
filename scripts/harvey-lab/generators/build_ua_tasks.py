#!/usr/bin/env python3
"""Build the two pilot Ukrainian LAB tasks (Track A diligence, Track B litigation).

Generates task.json plus a documents/ workspace of real-format .docx/.xlsx files.

Every party, code and address is synthetic. Ukrainian company codes (ЄДРПОУ) are
deliberately chosen to FAIL their checksum, so tests/test_no_real_identifiers.py
proves they cannot collide with a real registered entity. Matter structure is
modelled on public register shapes, not copied from any real matter.

Usage (on local.lex):
    uv run --with python-docx --with openpyxl python build_ua_tasks.py <tasks_root>
"""

import json
import sys
from pathlib import Path

from docx import Document
from docx.shared import Pt
from openpyxl import Workbook


# ── helpers ───────────────────────────────────────────────────────────


def _edrpou_checksum_ok(digits: str) -> bool:
    nums = [int(d) for d in digits]
    base = [1, 2, 3, 4, 5, 6, 7] if int(digits) < 30000000 else [7, 1, 2, 3, 4, 5, 6]
    c = sum(w * n for w, n in zip(base, nums[:7])) % 11
    if c >= 10:
        c = sum(w * n for w, n in zip([w + 2 for w in base], nums[:7])) % 11
        if c >= 10:
            return False
    return c == nums[7]


def assert_synthetic(code: str) -> str:
    """Fail loudly if a chosen code could belong to a real registered entity."""
    if _edrpou_checksum_ok(code):
        raise SystemExit(
            f"ЄДРПОУ {code} passes the checksum, so it may be a real company. "
            f"Pick a code that fails it."
        )
    return code


def uah(n: int) -> str:
    """Format an amount the Ukrainian way: 1250000 -> '1 250 000'.

    Formatting the NUMBER, never the whole sentence. Applying
    .replace(",", " ") to a formatted sentence also eats legitimate commas.
    """
    return f"{n:,}".replace(",", " ")


def doc(path: Path, title: str, blocks: list) -> None:
    """Write a .docx. blocks is a list of (style, text) where style is
    'h' heading, 'p' paragraph, 'b' bold paragraph."""
    d = Document()
    style = d.styles["Normal"]
    style.font.name = "Times New Roman"
    style.font.size = Pt(12)
    d.add_heading(title, level=1)
    for kind, text in blocks:
        if kind == "h":
            d.add_heading(text, level=2)
        elif kind == "b":
            p = d.add_paragraph()
            p.add_run(text).bold = True
        else:
            d.add_paragraph(text)
    path.parent.mkdir(parents=True, exist_ok=True)
    d.save(path)


def sheet(path: Path, title: str, header: list, rows: list) -> None:
    wb = Workbook()
    ws = wb.active
    ws.title = title[:31]
    ws.append(header)
    for r in rows:
        ws.append(r)
    for i, _ in enumerate(header, 1):
        ws.column_dimensions[ws.cell(row=1, column=i).column_letter].width = 26
    path.parent.mkdir(parents=True, exist_ok=True)
    wb.save(path)


def write_task(task_dir: Path, config: dict) -> None:
    task_dir.mkdir(parents=True, exist_ok=True)
    (task_dir / "task.json").write_text(
        json.dumps(config, ensure_ascii=False, indent=2) + "\n", encoding="utf-8"
    )


# ══════════════════════════════════════════════════════════════════════
# TRACK B — litigation: limitation period under martial law + Art. 625 CC
# ══════════════════════════════════════════════════════════════════════

# Synthetic parties. Codes verified checksum-invalid below.
PLAINTIFF = 'ТОВ "ВЕРБОВА ПРИСТАНЬ"'
PLAINTIFF_CODE = assert_synthetic("41234568")
DEFENDANT = 'ТОВ "КАМІНЕЦЬ ЛОГІСТИК"'
DEFENDANT_CODE = assert_synthetic("42345679")

DEBT = 1_250_000
CONTRACT_DATE = "15.03.2021"
DELIVERY_DATE = "02.04.2021"
DUE_DATE = "15.04.2021"
LIMITATION_START = "16.04.2021"
NAIVE_EXPIRY = "16.04.2024"
CLAIM_FILED = "20.09.2024"
JUDGMENT_DATE = "11.02.2025"


def build_litigation(root: Path) -> Path:
    t = root / "litigation-dispute-resolution" / "ua-limitation-period-martial-law"
    docs = t / "documents"

    doc(
        docs / "dogovir-postavky.docx",
        f"ДОГОВІР ПОСТАВКИ № 17/03-21 від {CONTRACT_DATE}",
        [
            ("p", f'Постачальник: {PLAINTIFF}, код ЄДРПОУ {PLAINTIFF_CODE}, '
                  f'м. Полтава, вул. Соборності, 14.'),
            ("p", f'Покупець: {DEFENDANT}, код ЄДРПОУ {DEFENDANT_CODE}, '
                  f'м. Дніпро, просп. Яворницького, 88.'),
            ("h", "1. Предмет договору"),
            ("p", "1.1. Постачальник зобов'язується поставити пиломатеріали "
                  "камерного сушіння в асортименті згідно зі специфікацією "
                  "(Додаток № 1), а Покупець - прийняти та оплатити їх."),
            ("h", "2. Ціна та порядок розрахунків"),
            ("p", f"2.1. Загальна вартість товару за цим Договором становить "
                  f"{uah(DEBT)} грн 00 коп. (один мільйон двісті п'ятдесят тисяч "
                  f"гривень) без ПДВ."),
            ("p", f"2.2. Покупець здійснює оплату протягом 13 (тринадцяти) "
                  f"календарних днів з дати підписання видаткової накладної."),
            ("h", "3. Відповідальність сторін"),
            ("p", "3.1. За порушення строків оплати Покупець сплачує пеню в "
                  "розмірі подвійної облікової ставки НБУ від суми "
                  "заборгованості за кожен день прострочення."),
            ("p", "3.2. Нарахування пені припиняється через шість місяців від "
                  "дня, коли зобов'язання мало бути виконано."),
            ("h", "4. Строк дії"),
            ("p", "4.1. Договір набирає чинності з моменту підписання і діє до "
                  "повного виконання сторонами своїх зобов'язань."),
        ],
    )

    doc(
        docs / "vydatkova-nakladna.docx",
        f"ВИДАТКОВА НАКЛАДНА № 204 від {DELIVERY_DATE}",
        [
            ("p", f"Постачальник: {PLAINTIFF}, код ЄДРПОУ {PLAINTIFF_CODE}"),
            ("p", f"Одержувач: {DEFENDANT}, код ЄДРПОУ {DEFENDANT_CODE}"),
            ("p", "Підстава: Договір поставки № 17/03-21 від 15.03.2021"),
            ("p", f"Пиломатеріали камерного сушіння, 250 куб. м, "
                  f"на загальну суму {uah(DEBT)} грн 00 коп."),
            ("p", "Товар прийнято без зауважень щодо кількості та якості."),
            ("p", "Від Покупця: комерційний директор Н. П. Гайдученко"),
        ],
    )

    doc(
        docs / "vymoga-pretenziya.docx",
        "ПРЕТЕНЗІЯ (вимога про оплату) від 12.05.2021",
        [
            ("p", f"Кому: {DEFENDANT}, код ЄДРПОУ {DEFENDANT_CODE}"),
            ("p", f"Від: {PLAINTIFF}, код ЄДРПОУ {PLAINTIFF_CODE}"),
            ("p", f"Товар за видатковою накладною № 204 від {DELIVERY_DATE} "
                  f"прийнято Вами без зауважень. Строк оплати сплив {DUE_DATE}."),
            ("p", f"Станом на дату цієї претензії заборгованість у сумі "
                  f"{uah(DEBT)} грн не погашена."),
            ("p", "Вимагаємо сплатити заборгованість протягом 7 календарних днів."),
            ("b", "Відповіді на претензію не надходило."),
        ],
    )

    doc(
        docs / "pozovna-zayava.docx",
        "ПОЗОВНА ЗАЯВА про стягнення заборгованості",
        [
            ("p", f"До Господарського суду Дніпропетровської області"),
            ("p", f"Позивач: {PLAINTIFF}, код ЄДРПОУ {PLAINTIFF_CODE}"),
            ("p", f"Відповідач: {DEFENDANT}, код ЄДРПОУ {DEFENDANT_CODE}"),
            ("p", f"Дата подання: {CLAIM_FILED}"),
            ("h", "Обставини справи"),
            ("p", f"Між сторонами укладено Договір поставки № 17/03-21 від "
                  f"{CONTRACT_DATE}. Товар поставлено {DELIVERY_DATE} згідно з "
                  f"видатковою накладною № 204. Оплата мала бути здійснена до "
                  f"{DUE_DATE} включно. Оплата не здійснена."),
            ("h", "Позовні вимоги"),
            ("p", f"1. Стягнути з Відповідача основну заборгованість "
                  f"{uah(DEBT)} грн 00 коп."),
            ("p", "2. Стягнути 3% річних відповідно до частини другої статті 625 "
                  "Цивільного кодексу України."),
            ("p", "3. Стягнути інфляційні втрати відповідно до частини другої "
                  "статті 625 Цивільного кодексу України."),
            ("p", "4. Стягнути пеню відповідно до пункту 3.1 Договору."),
        ],
    )

    doc(
        docs / "vidzyv-na-pozov.docx",
        "ВІДЗИВ НА ПОЗОВНУ ЗАЯВУ",
        [
            ("p", f"Відповідач: {DEFENDANT}, код ЄДРПОУ {DEFENDANT_CODE}"),
            ("b", "1. Сплив позовної давності"),
            ("p", f"Строк оплати сплив {DUE_DATE}. Перебіг загальної позовної "
                  f"давності розпочався {LIMITATION_START} і, за статтею 257 "
                  f"Цивільного кодексу України, мав завершитися {NAIVE_EXPIRY}. "
                  f"Позов подано {CLAIM_FILED}, тобто зі спливом позовної "
                  f"давності. Просимо відмовити в позові повністю."),
            ("b", "2. Щодо пені"),
            ("p", "Нарахування пені в будь-якому разі обмежене шістьма місяцями "
                  "згідно з пунктом 3.2 Договору."),
            ("p", "Заяву про застосування позовної давності зроблено до "
                  "ухвалення рішення у справі."),
        ],
    )

    doc(
        docs / "rishennya-sudu-pershoyi-instanciyi.docx",
        f"РІШЕННЯ Господарського суду Дніпропетровської області від {JUDGMENT_DATE}",
        [
            ("p", "Суд, дослідивши матеріали справи, встановив таке."),
            ("p", f"Факт поставки товару на суму {uah(DEBT)} грн підтверджено "
                  f"видатковою накладною № 204 від {DELIVERY_DATE}, підписаною "
                  f"уповноваженою особою Відповідача без зауважень."),
            ("p", f"Строк виконання грошового зобов'язання - {DUE_DATE}."),
            ("b", "Щодо заяви про застосування позовної давності"),
            ("p", "Суд відхиляє доводи Відповідача. Пунктом 19 Прикінцевих та "
                  "перехідних положень Цивільного кодексу України (в редакції "
                  "Закону № 3450-IX від 08.11.2023) встановлено, що у період "
                  "дії воєнного стану в Україні, введеного Указом Президента "
                  "України від 24 лютого 2022 року № 64/2022, перебіг позовної "
                  "давності, визначений цим Кодексом, зупиняється на строк дії "
                  "такого стану."),
            ("p", "Перебіг позовної давності розпочався 16.04.2021 і зупинився "
                  "24.02.2022, тобто до зупинення сплило 314 днів із трьох "
                  "років. Станом на дату подання позову воєнний стан не "
                  "припинено, отже перебіг давності не відновлювався."),
            ("p", "Отже, позовна давність не спливла, і позов подано в межах строку."),
            ("p", "Щодо вимоги про стягнення пені суд додатково зазначає, що до "
                  "вимог про стягнення неустойки (штрафу, пені) застосовується "
                  "спеціальна позовна давність в один рік (пункт 1 частини "
                  "другої статті 258 Цивільного кодексу України), перебіг якої "
                  "зупинено на тих самих підставах."),
            ("b", "Резолютивна частина"),
            ("p", f"Позов задовольнити частково. Стягнути з Відповідача основну "
                  f"заборгованість {uah(DEBT)} грн, 3% річних та інфляційні втрати. "
                  f"У стягненні пені відмовити з посиланням на пункт 3.2 "
                  f"Договору."),
        ],
    )

    sheet(
        docs / "rozrahunok-zaborhovanosti.xlsx",
        "Розрахунок",
        ["Показник", "Значення", "Підстава"],
        [
            ["Основна заборгованість, грн", DEBT, "Видаткова накладна № 204"],
            ["Дата виникнення прострочення", LIMITATION_START, "п. 2.2 Договору"],
            ["Ставка річних, %", 3, "ч. 2 ст. 625 ЦК України"],
            ["Період нарахування 3% річних", f"{LIMITATION_START} - {JUDGMENT_DATE}", "ч. 2 ст. 625 ЦК України"],
            ["Кількість днів прострочення", 1397, "розраховано"],
            ["Введення воєнного стану", "24.02.2022", "Указ Президента України"],
            ["Обмеження нарахування пені", "6 місяців", "п. 3.2 Договору"],
        ],
    )

    criteria = []

    def c(cid, title, match, source="expert", deliverable="analiz-perspektyv.docx"):
        criteria.append(
            {
                "id": cid,
                "title": title,
                "deliverables": [deliverable],
                "match_criteria": match,
                "source": source,
            }
        )

    c("C-001", "Identifies the general three-year limitation period",
      "PASS if the analysis states that the applicable general limitation period is three years. "
      "FAIL if it states any other duration or does not identify a limitation period at all.",
      source="oracle")
    c("C-002", "Cites Article 257 of the Civil Code for the general limitation period",
      "PASS if the analysis cites Article 257 of the Civil Code of Ukraine (стаття 257 ЦК України) "
      "as the source of the three-year general limitation period. FAIL if it cites a different "
      "article for that proposition or cites no provision.",
      source="oracle")
    c("C-003", "Fixes the start of the limitation period at 16.04.2021",
      "PASS if the analysis states that the limitation period began to run on 16.04.2021, the day "
      "after payment fell due on 15.04.2021. FAIL if it uses the contract date 15.03.2021, the "
      "delivery date 02.04.2021, or any other start date.",
      source="oracle")
    c("C-004", "Computes the unextended expiry as 16.04.2024",
      "PASS if the analysis states that, disregarding any extension, the limitation period would "
      "have expired on or about 16.04.2024. FAIL if it gives a materially different unextended "
      "expiry date.")
    c("C-005", "Identifies that the claim was filed after the unextended expiry",
      "PASS if the analysis notes that the claim was filed on 20.09.2024, which is after the "
      "unextended expiry of 16.04.2024. FAIL if it does not compare the filing date against the "
      "unextended expiry.")
    c("C-006", "Applies the martial-law SUSPENSION in the Civil Code transitional provisions",
      "PASS if the analysis identifies that the running of the limitation period is SUSPENDED "
      "(зупиняється) for the duration of martial law under paragraph 19 of the Final and "
      "Transitional Provisions of the Civil Code of Ukraine. FAIL if it does not identify the "
      "martial-law rule, if it describes the rule as an EXTENSION of the period (продовження) "
      "rather than a suspension of its running, or if it attributes the effect to the COVID-19 "
      "quarantine provision in paragraph 12, which extends periods but is a different rule.",
      source="oracle")
    c("C-006a", "Does not cite the quarantine article list for the martial-law rule",
      "PASS if the analysis does not attribute to paragraph 19 the list of articles "
      "'257, 258, 362, 559, 681, 728, 786, 1293'. That list belongs to paragraph 12 (COVID-19 "
      "quarantine); paragraph 19 suspends the running of limitation generally and contains no "
      "such list. FAIL if the analysis attaches that article list to the martial-law rule.",
      source="oracle")
    c("C-006b", "Computes the period that ran before suspension",
      "PASS if the analysis states that the limitation period ran from 16.04.2021 until martial "
      "law took effect on 24.02.2022, a period of approximately 314 days, and that the remainder "
      "of the three years has not yet resumed. FAIL if it treats the whole three years as having "
      "run, or does not separate the elapsed part from the suspended part.")
    c("C-007", "Dates the start of martial law to 24.02.2022",
      "PASS if the analysis states that martial law took effect on 24.02.2022. FAIL if it gives a "
      "different date or omits the date.",
      source="oracle")
    c("C-008", "Concludes the claim is NOT time-barred",
      "PASS if the analysis concludes that the limitation period had not expired when the claim "
      "was filed and that the defendant's limitation objection should fail. FAIL if it concludes "
      "the claim is time-barred, or leaves the conclusion open.")
    c("C-009", "Identifies Article 625(2) as the basis for 3% per annum and inflation losses",
      "PASS if the analysis cites part 2 of Article 625 of the Civil Code of Ukraine (частина "
      "друга статті 625 ЦК України) as the basis for both the 3% per annum and the inflation "
      "losses. FAIL if it cites a different provision or treats these as contractual claims.",
      source="oracle")
    c("C-010", "States the principal debt as 1 250 000 UAH",
      "PASS if the analysis states the principal debt as 1 250 000 UAH. FAIL if it states any "
      "other principal amount.",
      source="oracle")
    c("C-011", "Notes that the delivery note was signed without objection",
      "PASS if the analysis relies on the fact that consignment note No. 204 of 02.04.2021 was "
      "signed by the defendant's authorised representative without objection as to quantity or "
      "quality. FAIL if it does not address how delivery is evidenced.")
    c("C-012", "Identifies the contractual six-month cap on penalty",
      "PASS if the analysis identifies that clause 3.2 of the contract stops the accrual of "
      "penalty (пеня) six months after the obligation fell due. FAIL if it does not mention the "
      "six-month cap.")
    c("C-012a", "Identifies the one-year special limitation period for penalty claims",
      "PASS if the analysis identifies that claims for penalty (неустойка, штраф, пеня) are "
      "subject to a one-year special limitation period under point 1 of part 2 of Article 258 of "
      "the Civil Code of Ukraine, as distinct from the three-year general period. FAIL if it "
      "applies the three-year period to the penalty claim or does not address the special period.",
      source="oracle")
    c("C-012b", "Notes the one-year period is suspended on the same basis",
      "PASS if the analysis notes that the one-year period for the penalty claim is suspended by "
      "the same martial-law rule, so the penalty claim is not time-barred either and its refusal "
      "rests on the contractual cap rather than on limitation. FAIL if it concludes the penalty "
      "claim was refused because that one-year period expired.")
    c("C-013", "Explains why the penalty claim was refused",
      "PASS if the analysis explains that the first-instance court refused the penalty claim by "
      "reference to clause 3.2 of the contract. FAIL if it does not explain the refusal or "
      "attributes it to a different reason.")
    c("C-014", "Does not recommend appealing the part of the judgment already won",
      "PASS if the analysis limits any recommended appeal to the refused penalty claim and does "
      "not propose appealing the principal, 3% per annum, or inflation losses, which were "
      "granted. FAIL if it recommends appealing parts of the judgment decided in the client's "
      "favour.")
    c("C-015", "Distinguishes penalty (пеня) from Article 625 compensation",
      "PASS if the analysis explains that the 3% per annum and inflation losses under Article 625 "
      "are a separate remedy from the contractual penalty and are not subject to the clause 3.2 "
      "six-month cap. FAIL if it conflates the two or applies the cap to the Article 625 amounts.")
    c("C-016", "States the applicable per-annum rate as 3%",
      "PASS if the analysis states the statutory rate as 3% per annum. FAIL if it states a "
      "different rate.",
      source="oracle")
    c("C-017", "Identifies the correct first-instance court",
      "PASS if the analysis identifies the Commercial Court of Dnipropetrovsk Oblast "
      "(Господарський суд Дніпропетровської області) as the court of first instance. FAIL if it "
      "names a different court.")
    c("C-018", "Names the correct parties",
      "PASS if the analysis identifies ТОВ \"ВЕРБОВА ПРИСТАНЬ\" as claimant and "
      "ТОВ \"КАМІНЕЦЬ ЛОГІСТИК\" as defendant. FAIL if the roles are reversed or a party from "
      "outside this matter is named.")
    c("C-019", "Notes the limitation objection was raised in time",
      "PASS if the analysis notes that the defendant raised the limitation objection before "
      "judgment, so it was properly before the court. FAIL if it does not address the timing of "
      "the objection.")
    c("C-020", "Written in Ukrainian",
      "PASS if the deliverable is written in Ukrainian. FAIL if it is written in English, "
      "Russian, or any other language.")
    c("C-021", "Applies the law in force as at the judgment date",
      "PASS if the analysis applies paragraph 19 as a live provision, consistent with the "
      "instruction to apply the law in force on 11.02.2025. FAIL if it states that paragraph 19 "
      "does not exist or has been repealed. Paragraph 19 was in force from 15.03.2022 (Law "
      "2120-IX), restated by Law 3450-IX of 08.11.2023, and repealed only on 14.05.2025 by Law "
      "4434-IX, which is after the date this matter is set on.",
      source="oracle")

    write_task(
        t,
        {
            "title": "Аналіз перспектив апеляції: позовна давність під час воєнного стану",
            "work_type": "analyze",
            "language": "uk",
            "jurisdiction": "UA",
            "judge_language": "en",
            "tags": [
                "Litigation (General)", "limitation-period", "martial-law",
                "civil-code-ukraine", "art-625", "supply-contract", "debt-recovery",
                "ukraine",
            ],
            "instructions": (
                "Ознайомтеся з матеріалами справи та підготуйте для клієнта "
                "(позивача) аналіз перспектив апеляційного оскарження в частині, "
                "у якій суд відмовив. Оцініть також ризик скасування рішення за "
                "доводом відповідача про сплив позовної давності. "
                "Застосовуйте законодавство в редакції, чинній станом на дату "
                "ухвалення рішення суду першої інстанції (11.02.2025). "
                "Результат: `analiz-perspektyv.docx`."
            ),
            "deliverables": {"analiz-perspektyv.docx": "analiz-perspektyv.docx"},
            "criteria": criteria,
        },
    )
    return t


# ══════════════════════════════════════════════════════════════════════
# TRACK A — counterparty diligence across Ukrainian public registers
# ══════════════════════════════════════════════════════════════════════

TARGET = 'ТОВ "ГРАНІТ-ІНВЕСТ ЮА"'
TARGET_CODE = assert_synthetic("43456780")
PARENT = 'ТОВ "ПІВДЕННА ХОЛДИНГОВА ГРУПА"'
PARENT_CODE = assert_synthetic("44567891")
SANCTIONED = 'АТ "ТРАНСБУД РЕСУРС"'
SANCTIONED_CODE = assert_synthetic("45678902")
CLIENT = 'ТОВ "АГРОТЕХ ПАРТНЕРС"'


def build_diligence(root: Path) -> Path:
    t = root / "diligence" / "ua-counterparty-register-screening"
    docs = t / "documents"

    doc(
        docs / "zapyt-kliyenta.docx",
        "ЗАПИТ КЛІЄНТА",
        [
            ("p", f"Від: {CLIENT}, фінансовий директор"),
            ("p", "Ми плануємо укласти рамковий договір поставки з "
                  f"{TARGET} з відстрочкою платежу 90 днів та авансом 30% "
                  f"(орієнтовно 4 200 000 грн)."),
            ("p", "Просимо перевірити контрагента за відкритими реєстрами та "
                  "повідомити, чи є підстави відмовитися від угоди або "
                  "змінити її умови."),
            ("p", "Наша внутрішня політика забороняє передоплату контрагентам, "
                  "щодо яких відкрито провадження у справі про банкрутство."),
        ],
    )

    doc(
        docs / "vytyag-EDR.docx",
        "ВИТЯГ З ЄДИНОГО ДЕРЖАВНОГО РЕЄСТРУ ЮРИДИЧНИХ ОСІБ",
        [
            ("p", f"Найменування: {TARGET}"),
            ("p", f"Код ЄДРПОУ: {TARGET_CODE}"),
            ("p", "Організаційно-правова форма: ТОВАРИСТВО З ОБМЕЖЕНОЮ ВІДПОВІДАЛЬНІСТЮ"),
            ("p", "Дата державної реєстрації: 11.06.2019"),
            ("b", "Стан: в стані припинення"),
            ("p", "Розмір статутного капіталу: 100 000,00 грн"),
            ("p", "Місцезнаходження: м. Кривий Ріг, вул. Каштанова, 3, оф. 210"),
            ("h", "Засновники (учасники)"),
            ("p", f"{PARENT}, код ЄДРПОУ {PARENT_CODE}; розмір частки - 100%"),
            ("h", "Кінцевий бенефіціарний власник"),
            ("p", "ЗАЛІСНИЙ ОСТАП ВАЛЕРІЙОВИЧ; тип впливу - вирішальний вплив; "
                  "відсоток частки - 100"),
            ("h", "Керівник"),
            ("p", "ЗАЛІСНИЙ ОСТАП ВАЛЕРІЙОВИЧ"),
        ],
    )

    doc(
        docs / "vytyag-EDR-zasnovnyk.docx",
        "ВИТЯГ З ЄДИНОГО ДЕРЖАВНОГО РЕЄСТРУ ЮРИДИЧНИХ ОСІБ (засновник)",
        [
            ("p", f"Найменування: {PARENT}"),
            ("p", f"Код ЄДРПОУ: {PARENT_CODE}"),
            ("p", "Стан: зареєстровано"),
            ("h", "Кінцевий бенефіціарний власник"),
            ("p", "ЗАЛІСНИЙ ОСТАП ВАЛЕРІЙОВИЧ; тип впливу - вирішальний вплив; "
                  "відсоток частки - 100"),
            ("h", "Інші юридичні особи, де ця особа є кінцевим бенефіціарним власником"),
            ("p", f"{SANCTIONED}, код ЄДРПОУ {SANCTIONED_CODE}; відсоток частки - 75"),
        ],
    )

    sheet(
        docs / "vykonavchi-provadzhennya.xlsx",
        "Виконавчі провадження",
        ["Номер ВП", "Боржник", "Код ЄДРПОУ", "Стягувач", "Сума, грн", "Дата відкриття", "Стан"],
        [
            ["71204588", TARGET, TARGET_CODE, 'ТОВ "ДНІПРОЕНЕРГОЗБУТ"', 1840000, "14.03.2024", "відкрито"],
            ["71880341", TARGET, TARGET_CODE, "ГУ ДПС у Дніпропетровській обл.", 962400, "07.06.2024", "відкрито"],
            ["69551027", TARGET, TARGET_CODE, 'ТОВ "СПЕЦМОНТАЖ ПЛЮС"', 145000, "22.08.2023", "завершено"],
        ],
    )

    sheet(
        docs / "sudovi-spravy.xlsx",
        "Судові справи",
        ["Номер справи", "Суд", "Роль", "Предмет", "Сума позову, грн", "Дата"],
        [
            ["904/1188/24", "Господарський суд Дніпропетровської області", "відповідач", "стягнення заборгованості", 1840000, "05.02.2024"],
            ["904/2447/24", "Господарський суд Дніпропетровської області", "відповідач", "справа про банкрутство", 0, "18.07.2024"],
            ["904/0912/23", "Господарський суд Дніпропетровської області", "позивач", "стягнення збитків", 320000, "14.04.2023"],
        ],
    )

    doc(
        docs / "dovidka-podatkovyi-borh.docx",
        "ДОВІДКА ПРО СТАН РОЗРАХУНКІВ З БЮДЖЕТОМ",
        [
            ("p", f"Платник: {TARGET}, код ЄДРПОУ {TARGET_CODE}"),
            ("p", "Станом на 01.08.2024"),
            ("b", "Загальна сума податкового боргу: 962 400,00 грн"),
            ("p", "У тому числі: податок на прибуток - 611 200,00 грн; "
                  "ПДВ - 351 200,00 грн."),
            ("p", "Платник включений до переліку боржників, оприлюдненого "
                  "відповідно до законодавства."),
        ],
    )

    doc(
        docs / "sankciyni-perevirky.docx",
        "РЕЗУЛЬТАТИ ПЕРЕВІРКИ ЗА САНКЦІЙНИМИ ПЕРЕЛІКАМИ",
        [
            ("p", f"{TARGET}, код ЄДРПОУ {TARGET_CODE}: у переліках не виявлено."),
            ("p", f"{PARENT}, код ЄДРПОУ {PARENT_CODE}: у переліках не виявлено."),
            ("b", f"{SANCTIONED}, код ЄДРПОУ {SANCTIONED_CODE}: включено до "
                  f"переліку осіб, до яких застосовано санкції, рішенням від "
                  f"19.10.2023. Строк застосування - 10 років."),
            ("p", "ЗАЛІСНИЙ ОСТАП ВАЛЕРІЙОВИЧ: як фізична особа у переліках "
                  "не виявлений."),
        ],
    )

    sheet(
        docs / "arma-aktyvy.xlsx",
        "Арештовані активи",
        ["Номер справи", "Власник", "Код ЄДРПОУ", "Тип активу", "Опис", "Дата арешту"],
        [
            ["904/2447/24", TARGET, TARGET_CODE, "нерухоме майно",
             "виробничий комплекс, м. Кривий Ріг, вул. Каштанова, 3", "02.08.2024"],
        ],
    )

    criteria = []

    def c(cid, title, match, source="expert", deliverable="dyu-dilidzhens-memo.docx"):
        criteria.append(
            {
                "id": cid,
                "title": title,
                "deliverables": [deliverable],
                "match_criteria": match,
                "source": source,
            }
        )

    c("C-001", "Flags that the counterparty is in the process of termination",
      "PASS if the memo flags that the register shows the counterparty's status as "
      "'в стані припинення' (termination in progress). FAIL if it reports the counterparty as "
      "actively registered or omits the status entirely.",
      source="oracle")
    c("C-002", "Reports the correct company code for the counterparty",
      f"PASS if the memo identifies the counterparty's ЄДРПОУ code as {TARGET_CODE}. FAIL if it "
      f"reports a different code or omits it.",
      source="oracle")
    c("C-003", "Identifies an open bankruptcy case",
      "PASS if the memo identifies that case 904/2447/24, opened 18.07.2024, is a bankruptcy case "
      "in which the counterparty is respondent. FAIL if it does not identify the bankruptcy case "
      "or misclassifies it as an ordinary debt claim.",
      source="oracle")
    c("C-004", "Applies the client's stated policy against prepayment",
      "PASS if the memo states that the client's internal policy prohibits prepayment to a "
      "counterparty subject to bankruptcy proceedings, and that the proposed 30% advance is "
      "therefore not permitted. FAIL if it does not connect the bankruptcy finding to the client's "
      "stated policy.")
    c("C-005", "Recommends against the advance payment",
      "PASS if the memo recommends against paying the 30% advance, or conditions it on security. "
      "FAIL if it recommends proceeding with the advance as proposed.")
    c("C-006", "Quantifies the open enforcement proceedings",
      "PASS if the memo states the total value of the two OPEN enforcement proceedings as "
      "2 802 400 UAH (1 840 000 plus 962 400). FAIL if it reports a different total, or includes "
      "the completed proceeding 69551027 in the total.",
      source="oracle")
    c("C-007", "Excludes the completed enforcement proceeding from the exposure figure",
      "PASS if the memo treats proceeding 69551027 (145 000 UAH) as completed and excludes it "
      "from current exposure. FAIL if it counts it as open.",
      source="oracle")
    c("C-008", "Reports the tax debt",
      "PASS if the memo reports the counterparty's tax debt as 962 400 UAH as at 01.08.2024. "
      "FAIL if it reports a different figure or omits the tax debt.",
      source="oracle")
    c("C-009", "Links the tax debt to the enforcement proceeding",
      "PASS if the memo notes that enforcement proceeding 71880341 (962 400 UAH, creditor the "
      "tax authority) corresponds to the same tax debt, so the two figures are not additive. "
      "FAIL if it adds the tax debt to that enforcement proceeding as separate exposure.")
    c("C-010", "Traces ownership to the parent company",
      f"PASS if the memo states that {PARENT} holds 100% of the counterparty. FAIL if it reports a "
      f"different shareholder or a different percentage.",
      source="oracle")
    c("C-011", "Identifies the ultimate beneficial owner",
      "PASS if the memo identifies ЗАЛІСНИЙ ОСТАП ВАЛЕРІЙОВИЧ as the ultimate beneficial owner "
      "with decisive influence and a 100% interest. FAIL if it names a different individual or "
      "omits the UBO.",
      source="oracle")
    c("C-012", "Connects the UBO to the sanctioned entity",
      f"PASS if the memo identifies that the same ultimate beneficial owner holds 75% of "
      f"{SANCTIONED}, which is subject to sanctions. FAIL if it does not make this connection, "
      f"which requires reading the shareholder's register extract together with the sanctions "
      f"screening.",
      source="oracle")
    c("C-013", "States that the counterparty itself is not sanctioned",
      "PASS if the memo states that neither the counterparty nor its direct shareholder appears "
      "on a sanctions list, while still flagging the UBO connection. FAIL if it asserts that the "
      "counterparty itself is sanctioned.")
    c("C-014", "Dates the sanctions decision",
      "PASS if the memo states that sanctions were imposed on the connected entity by a decision "
      "of 19.10.2023 for a term of 10 years. FAIL if it gives a different date or term.",
      source="oracle")
    c("C-015", "Identifies the seized production complex",
      "PASS if the memo identifies that the counterparty's production complex at "
      "вул. Каштанова, 3, м. Кривий Ріг was seized on 02.08.2024 in case 904/2447/24. FAIL if it "
      "does not report the asset seizure.",
      source="oracle")
    c("C-016", "Notes that the seized asset is the registered address",
      "PASS if the memo notes that the seized production complex is at the same address as the "
      "counterparty's registered office (вул. Каштанова, 3), so its operating base is encumbered. "
      "FAIL if it does not connect the two.")
    c("C-017", "Distinguishes the counterparty's own claim from claims against it",
      "PASS if the memo treats case 904/0912/23 as one where the counterparty is claimant, and "
      "does not count it as exposure. FAIL if it counts that case as a claim against the "
      "counterparty.",
      source="oracle")
    c("C-018", "Gives an overall recommendation against contracting on the proposed terms",
      "PASS if the memo concludes that the counterparty should not be engaged on the proposed "
      "terms. FAIL if it recommends proceeding as proposed, or gives no clear recommendation.")
    c("C-019", "Does not invent findings absent from the documents",
      "PASS if every adverse finding in the memo traces to a supplied document. FAIL if the memo "
      "asserts criminal proceedings, licence revocations, or other adverse facts that appear "
      "nowhere in the workspace.")
    c("C-020", "Written in Ukrainian",
      "PASS if the deliverable is written in Ukrainian. FAIL if it is written in English, "
      "Russian, or any other language.")

    write_task(
        t,
        {
            "title": "Перевірка контрагента за відкритими реєстрами перед укладенням договору",
            "work_type": "analyze",
            "language": "uk",
            "jurisdiction": "UA",
            "judge_language": "en",
            "tags": [
                "Diligence", "counterparty-screening", "beneficial-ownership",
                "sanctions", "bankruptcy", "enforcement-proceedings", "ukraine",
            ],
            "instructions": (
                "Клієнт розглядає укладення рамкового договору поставки з "
                "контрагентом. Перевірте контрагента за наданими витягами та "
                "реєстровими вивантаженнями і підготуйте меморандум про "
                "виявлені ризики з чіткою рекомендацією. "
                "Результат: `dyu-dilidzhens-memo.docx`."
            ),
            "deliverables": {"dyu-dilidzhens-memo.docx": "dyu-dilidzhens-memo.docx"},
            "criteria": criteria,
        },
    )
    return t


def main():
    root = Path(sys.argv[1]) if len(sys.argv) > 1 else Path("tasks")
    a = build_diligence(root)
    b = build_litigation(root)
    for t in (a, b):
        cfg = json.loads((t / "task.json").read_text(encoding="utf-8"))
        n_oracle = sum(1 for c in cfg["criteria"] if c.get("source") == "oracle")
        n_docs = len(list((t / "documents").iterdir()))
        print(f"{t}: {len(cfg['criteria'])} criteria "
              f"({n_oracle} oracle), {n_docs} documents")


if __name__ == "__main__":
    main()
