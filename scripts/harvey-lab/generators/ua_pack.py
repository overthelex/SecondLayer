#!/usr/bin/env python3
"""Ukrainian LAB pack: register-screening diligence family.

Mirrors how LAB already organises `diligence` (eight structurally similar
merger/buyout reviews with different facts). Each scenario here supplies its own
entities, register extracts and traps; the rubric is generated from those facts,
so every criterion names concrete values rather than a generic expectation.

Design rule: the truth of every criterion is contained in the supplied
documents. No criterion depends on a statutory provision, so correctness is
checkable by internal consistency rather than by legal research. Statutory
tasks live in the litigation family and reuse only provisions verified against
zakon.rada.gov.ua.

Usage:
    uv run --with python-docx --with openpyxl python ua_pack.py <tasks_root>
"""

import json
import sys
from dataclasses import dataclass, field
from pathlib import Path

from docx import Document
from docx.shared import Pt
from openpyxl import Workbook


# ── synthetic-code helpers ────────────────────────────────────────────


def _edrpou_ok(d: str) -> bool:
    n = [int(c) for c in d]
    base = [1, 2, 3, 4, 5, 6, 7] if int(d) < 30000000 else [7, 1, 2, 3, 4, 5, 6]
    c = sum(w * x for w, x in zip(base, n[:7])) % 11
    if c >= 10:
        c = sum(w * x for w, x in zip([w + 2 for w in base], n[:7])) % 11
        if c >= 10:
            return False
    return c == n[7]


def synth(code: str) -> str:
    """A code that must NOT pass the checksum, so it cannot be a real company."""
    if _edrpou_ok(code):
        raise SystemExit(f"{code} passes the ЄДРПОУ checksum; pick another")
    return code


def uah(n: int) -> str:
    return f"{n:,}".replace(",", " ")


# ── rendering ─────────────────────────────────────────────────────────


def doc(path: Path, title: str, blocks) -> None:
    d = Document()
    st = d.styles["Normal"]
    st.font.name = "Times New Roman"
    st.font.size = Pt(12)
    d.add_heading(title, level=1)
    for kind, text in blocks:
        if kind == "h":
            d.add_heading(text, level=2)
        elif kind == "b":
            d.add_paragraph().add_run(text).bold = True
        else:
            d.add_paragraph(text)
    path.parent.mkdir(parents=True, exist_ok=True)
    d.save(path)


def sheet(path: Path, tab: str, header, rows) -> None:
    wb = Workbook()
    ws = wb.active
    ws.title = tab[:31]
    ws.append(header)
    for r in rows:
        ws.append(r)
    for i in range(1, len(header) + 1):
        ws.column_dimensions[ws.cell(row=1, column=i).column_letter].width = 26
    path.parent.mkdir(parents=True, exist_ok=True)
    wb.save(path)


# ── scenario model ────────────────────────────────────────────────────


@dataclass
class Proceeding:
    number: str
    creditor: str
    amount: int
    opened: str
    state: str  # "відкрито" | "завершено" | "повернуто стягувачу"


@dataclass
class CourtCase:
    number: str
    role: str  # "відповідач" | "позивач"
    subject: str
    amount: int
    date: str


@dataclass
class Scenario:
    slug: str
    title: str
    industry: str
    deal: str
    target: str
    target_code: str
    status: str
    registered: str
    address: str
    capital: int
    parent: str
    parent_code: str
    parent_share: int
    ubo: str
    director: str
    proceedings: list
    cases: list
    tax_debt: int
    tax_date: str
    client_policy: str
    policy_trigger: str
    hidden_link: str
    hidden_detail: str
    deliverable: str = "dyu-dilidzhens-memo.docx"
    extra_docs: list = field(default_factory=list)
    extra_criteria: list = field(default_factory=list)

    @property
    def open_total(self) -> int:
        return sum(p.amount for p in self.proceedings if p.state == "відкрито")

    @property
    def closed(self) -> list:
        return [p for p in self.proceedings if p.state != "відкрито"]

    @property
    def as_claimant(self) -> list:
        return [c for c in self.cases if c.role == "позивач"]


# ── document generation ───────────────────────────────────────────────


def build_documents(s: Scenario, docs: Path) -> None:
    doc(docs / "zapyt-kliyenta.docx", "ЗАПИТ КЛІЄНТА", [
        ("p", f"Галузь: {s.industry}."),
        ("p", f"Ми плануємо {s.deal} з контрагентом {s.target}."),
        ("p", "Просимо перевірити контрагента за наданими витягами та "
              "реєстровими вивантаженнями і повідомити про виявлені ризики."),
        ("b", f"Внутрішня політика: {s.client_policy}"),
    ])

    doc(docs / "vytyag-EDR.docx",
        "ВИТЯГ З ЄДИНОГО ДЕРЖАВНОГО РЕЄСТРУ ЮРИДИЧНИХ ОСІБ", [
            ("p", f"Найменування: {s.target}"),
            ("p", f"Код ЄДРПОУ: {s.target_code}"),
            ("p", "Організаційно-правова форма: ТОВАРИСТВО З ОБМЕЖЕНОЮ "
                  "ВІДПОВІДАЛЬНІСТЮ"),
            ("p", f"Дата державної реєстрації: {s.registered}"),
            ("b", f"Стан: {s.status}"),
            ("p", f"Розмір статутного капіталу: {uah(s.capital)},00 грн"),
            ("p", f"Місцезнаходження: {s.address}"),
            ("h", "Засновники (учасники)"),
            ("p", f"{s.parent}, код ЄДРПОУ {s.parent_code}; "
                  f"розмір частки - {s.parent_share}%"),
            ("h", "Кінцевий бенефіціарний власник"),
            ("p", f"{s.ubo}; тип впливу - вирішальний вплив; "
                  f"відсоток частки - {s.parent_share}"),
            ("h", "Керівник"),
            ("p", s.director),
        ])

    doc(docs / "vytyag-EDR-zasnovnyk.docx",
        "ВИТЯГ З ЄДР (засновник контрагента)", [
            ("p", f"Найменування: {s.parent}"),
            ("p", f"Код ЄДРПОУ: {s.parent_code}"),
            ("p", "Стан: зареєстровано"),
            ("h", "Кінцевий бенефіціарний власник"),
            ("p", f"{s.ubo}; тип впливу - вирішальний вплив; відсоток частки - 100"),
            ("h", "Інші зв'язки цієї особи за даними реєстру"),
            ("b", s.hidden_detail),
        ])

    sheet(docs / "vykonavchi-provadzhennya.xlsx", "Виконавчі провадження",
          ["Номер ВП", "Боржник", "Код ЄДРПОУ", "Стягувач", "Сума, грн",
           "Дата відкриття", "Стан"],
          [[p.number, s.target, s.target_code, p.creditor, p.amount,
            p.opened, p.state] for p in s.proceedings])

    sheet(docs / "sudovi-spravy.xlsx", "Судові справи",
          ["Номер справи", "Суд", "Роль контрагента", "Предмет",
           "Сума, грн", "Дата"],
          [[c.number, "Господарський суд", c.role, c.subject, c.amount, c.date]
           for c in s.cases])

    doc(docs / "dovidka-podatkovyi-borh.docx",
        "ДОВІДКА ПРО СТАН РОЗРАХУНКІВ З БЮДЖЕТОМ", [
            ("p", f"Платник: {s.target}, код ЄДРПОУ {s.target_code}"),
            ("p", f"Станом на {s.tax_date}"),
            ("b", f"Загальна сума податкового боргу: {uah(s.tax_debt)},00 грн"),
        ])

    for fname, title, blocks in s.extra_docs:
        doc(docs / fname, title, blocks)


# ── rubric generation ─────────────────────────────────────────────────


def build_criteria(s: Scenario) -> list:
    out, n = [], 0

    def add(title, match, source="expert"):
        nonlocal n
        n += 1
        out.append({
            "id": f"C-{n:03d}", "title": title,
            "deliverables": [s.deliverable],
            "match_criteria": match, "source": source,
        })

    add("Reports the counterparty's registered status",
        f"PASS if the memo reports the counterparty's register status as "
        f"'{s.status}'. FAIL if it reports a different status or omits it.",
        "oracle")
    add("Reports the correct company code",
        f"PASS if the memo identifies the counterparty's ЄДРПОУ code as "
        f"{s.target_code}. FAIL if it gives a different code or omits it.",
        "oracle")
    add("Quantifies open enforcement exposure",
        f"PASS if the memo states the total of the OPEN enforcement "
        f"proceedings as {uah(s.open_total)} UAH. FAIL if it reports a "
        f"different total.", "oracle")
    if s.closed:
        names = ", ".join(p.number for p in s.closed)
        add("Excludes non-open enforcement proceedings from exposure",
            f"PASS if the memo treats proceeding(s) {names} as not currently "
            f"open and excludes them from the exposure total. FAIL if it "
            f"counts them as open.", "oracle")
    add("Reports the tax debt",
        f"PASS if the memo reports the tax debt as {uah(s.tax_debt)} UAH as at "
        f"{s.tax_date}. FAIL if it gives a different figure or omits it.",
        "oracle")
    add("Traces ownership to the direct shareholder",
        f"PASS if the memo states that {s.parent} holds {s.parent_share}% of "
        f"the counterparty. FAIL if it names a different shareholder or share.",
        "oracle")
    add("Identifies the ultimate beneficial owner",
        f"PASS if the memo identifies {s.ubo} as the ultimate beneficial "
        f"owner. FAIL if it names a different individual or omits the UBO.",
        "oracle")
    add("Surfaces the connection that requires joining two documents",
        f"PASS if the memo identifies the following, which is only visible by "
        f"reading the shareholder's register extract together with the other "
        f"documents: {s.hidden_link} FAIL if the memo does not make this "
        f"connection.", "oracle")
    if s.as_claimant:
        names = ", ".join(c.number for c in s.as_claimant)
        add("Does not count the counterparty's own claims as exposure",
            f"PASS if the memo treats case(s) {names} as brought BY the "
            f"counterparty and excludes them from exposure against it. FAIL if "
            f"it counts them as claims against the counterparty.", "oracle")
    add("Applies the client's stated policy",
        f"PASS if the memo applies the client's stated policy ({s.client_policy}) "
        f"to the facts and reaches the conclusion it compels, namely that "
        f"{s.policy_trigger}. FAIL if the memo does not connect its findings to "
        f"the stated policy.")
    add("Gives a clear recommendation",
        "PASS if the memo ends with an unambiguous recommendation on whether to "
        "proceed, and on what conditions. FAIL if it lists risks without "
        "reaching a recommendation.")
    add("Does not invent findings absent from the documents",
        "PASS if every adverse finding traces to a supplied document. FAIL if "
        "the memo asserts criminal proceedings, licence revocations, "
        "insolvency or other adverse facts that appear nowhere in the "
        "workspace.")
    add("Reports the registered address",
        f"PASS if the memo states the counterparty's registered address as "
        f"{s.address}. FAIL if it gives a different address.", "oracle")
    add("Reports the share capital",
        f"PASS if the memo states the share capital as {uah(s.capital)} UAH. "
        f"FAIL if it gives a different figure.", "oracle")
    add("Written in Ukrainian",
        "PASS if the deliverable is written in Ukrainian. FAIL if it is in "
        "English, Russian or any other language.")

    for title, match, source in s.extra_criteria:
        add(title, match, source)
    return out


def build_task(s: Scenario, root: Path) -> Path:
    t = root / "diligence" / s.slug
    build_documents(s, t / "documents")
    cfg = {
        "title": s.title,
        "work_type": "analyze",
        "language": "uk",
        "jurisdiction": "UA",
        "judge_language": "en",
        "tags": ["Diligence", "counterparty-screening", "ukraine",
                 "beneficial-ownership", "enforcement-proceedings"],
        "instructions": (
            f"Клієнт розглядає {s.deal} з контрагентом. Перевірте контрагента "
            f"за наданими витягами та реєстровими вивантаженнями і підготуйте "
            f"меморандум про виявлені ризики з чіткою рекомендацією. "
            f"Результат: `{s.deliverable}`."
        ),
        "deliverables": {s.deliverable: s.deliverable},
        "criteria": build_criteria(s),
    }
    t.mkdir(parents=True, exist_ok=True)
    (t / "task.json").write_text(
        json.dumps(cfg, ensure_ascii=False, indent=2) + "\n", encoding="utf-8"
    )
    return t
