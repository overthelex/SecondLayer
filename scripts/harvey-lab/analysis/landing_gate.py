#!/usr/bin/env python3
"""Deterministic deliverable landing gate, applied after the fact.

Reimplements what Niklaus' `deliverable_landing_gate` does inside the harness:
put each requested deliverable at the top level of output/ under the exact
requested filename. Here it runs as a post-hoc pass over completed runs, so the
same rollouts can be judged twice and the difference attributed to landing
alone rather than to any change in the model's work.

Selection is deliberately dumb and explainable:
  1. exact requested name already present -> nothing to do
  2. same stem, any extension            -> that file
  3. otherwise the largest text-bearing file in output/, excluding scripts

.docx is produced from markdown/text by wrapping the text in a minimal Word
document, because the judge reads the file, not its styling.

Usage:
    uv run --with python-docx python landing_gate.py <results_root> [--apply]
"""

import json
import shutil
import sys
from pathlib import Path

SCRIPTY = {".py", ".sh", ".json", ".log"}

# A file only counts as a deliverable if it carries real content. The pilot
# showed the agent writing its work to a working name and leaving a stub at the
# requested one: an 11-byte "placeholder" .docx, several 4-byte .md files. A
# gate that treats those as landed would report success and score zero.
MIN_CHARS = 200


def text_of(p: Path) -> str:
    if p.suffix.lower() == ".docx":
        try:
            import docx
            return "\n".join(x.text for x in docx.Document(p).paragraphs)
        except Exception:
            return ""
    try:
        return p.read_text(encoding="utf-8", errors="ignore")
    except Exception:
        return ""


def pick(out_dir: Path, wanted: str):
    exact = out_dir / wanted
    if exact.exists() and len(text_of(exact).strip()) >= MIN_CHARS:
        return None, "already-landed"

    substantive = [p for p in out_dir.rglob("*")
                   if p.is_file()
                   and p.suffix.lower() not in SCRIPTY
                   and len(text_of(p).strip()) >= MIN_CHARS]
    if not substantive:
        return None, "nothing-to-land"

    stem = Path(wanted).stem
    same_stem = [p for p in substantive if p.stem == stem]
    if same_stem:
        return max(same_stem, key=lambda p: len(text_of(p))), "same-stem"
    return max(substantive, key=lambda p: len(text_of(p))), "largest-text"


def land(src: Path, dst: Path) -> None:
    if dst.suffix.lower() == ".docx" and src.suffix.lower() != ".docx":
        from docx import Document
        d = Document()
        for line in text_of(src).splitlines():
            d.add_paragraph(line)
        d.save(dst)
    else:
        shutil.copy2(src, dst)


def main():
    root = Path(sys.argv[1])
    apply = "--apply" in sys.argv
    stats = {}
    for cfg_path in sorted(root.rglob("config.json")):
        run = cfg_path.parent
        out = run / "output"
        if not out.is_dir():
            continue
        cfg = json.loads(cfg_path.read_text(encoding="utf-8"))
        task = cfg.get("task") or cfg.get("task_id")
        wanted = list((cfg.get("deliverables") or {}).keys())
        if not wanted:
            tj = Path("tasks") / task / "task.json"
            if tj.exists():
                wanted = list(json.loads(tj.read_text(encoding="utf-8"))
                              .get("deliverables", {}).keys())
        for w in wanted:
            src, why = pick(out, w)
            stats[why] = stats.get(why, 0) + 1
            if src and apply:
                land(src, out / w)
            print(f"{why:16s} {task:58s} {src.name if src else '-':28s} -> {w}")
    print("\n" + json.dumps(stats, indent=1))


if __name__ == "__main__":
    main()
